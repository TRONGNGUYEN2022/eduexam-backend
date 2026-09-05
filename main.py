# -*- coding: utf-8 -*-
import os
import io
import json
import re
import random
import base64
import sqlite3
import subprocess
import requests
from typing import Optional, List, Dict, Any

from fastapi import FastAPI, HTTPException, Query, Request, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, HTMLResponse
from pydantic import BaseModel

import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
import pypdf

try:
    import pymupdf as fitz
except ImportError:
    try:
        import fitz
    except ImportError:
        fitz = None

try:
    from pdf2docx import Converter
except ImportError:
    Converter = None

try:
    import lxml.etree as ET
    from latex2mathml.converter import convert as latex_to_mathml
except ImportError:
    ET = None
    latex_to_mathml = None

from google import genai
from google.genai import types
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload

app = FastAPI(title="EduExam Multi-Level Pro Backend API - GDPT 2018 Engine", version="20.1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["*"],
)

DATA_DIR = "/var/data" if os.path.exists("/var/data") else "."
DB_FILE = os.path.join(DATA_DIR, "exam_bank.db")
CACHE_VAULT_DIR = os.path.join(DATA_DIR, "exam_vault")
os.makedirs(CACHE_VAULT_DIR, exist_ok=True)
DB_DRIVE_FILE_NAME = "exam_bank.db"
MOCK_TEX_FILE = "mock_exam_gdpt2018.tex"

DEFAULT_ENV_KEY = os.getenv("GEMINI_API_KEY", "")
global_ai_client = genai.Client(api_key=DEFAULT_ENV_KEY) if DEFAULT_ENV_KEY else None

GOOGLE_DRIVE_FOLDER_ID = os.getenv("GOOGLE_DRIVE_ROOT_FOLDER_ID", "")
SERVICE_ACCOUNT_FILE = os.getenv("GOOGLE_SERVICE_ACCOUNT_FILE", "service_account.json")
SERVICE_ACCOUNT_INFO_JSON = os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON", "")

# =========================================================================
# 1. DATA SCHEMAS (ĐƯA LÊN TRƯỚC ĐỂ TRÁNH LỖI NAMERROR)
# =========================================================================
class GenerateExamRequest(BaseModel):
    level: str = "thcs"
    subject: str = "Toán học"
    grade: int = 7
    topic: str = "Hình học trực quan"
    lesson: Optional[str] = "Bài học chung"
    cognitive_level: Optional[str] = "M2"
    duration_minutes: int = 45
    exam_title: str = "ĐỀ KIỂM TRA ĐỊNH KỲ"
    matrix_spec: Optional[str] = None
    source_notes: Optional[str] = None
    essay_mode: str = "ai_auto"
    essay_teacher_answer: Optional[str] = None
    selected_question_ids: Optional[List[int]] = []
    model_name: str = "gemini-2.5-flash"
    api_key: Optional[str] = None

class QuestionSave(BaseModel):
    level_stage: str = "thcs"
    grade: int
    subject: str = "Toán học"
    topic: str
    lesson: Optional[str] = "Bài học chung"
    level: str = "Thông hiểu"
    cognitive_level: Optional[str] = "M2"
    q_type: str
    content: str
    options: Optional[List[str]] = []
    correct_answer: Optional[str] = ""
    explanation: Optional[str] = ""
    has_image: Optional[int] = 0
    image_base64: Optional[str] = ""
    tikz_code: Optional[str] = ""

class DeficitRequest(BaseModel):
    level_stage: str = "thcs"
    grade: int
    subject: str = "Toán học"
    topic: str
    level: str
    q_type: str
    needed_count: int = 2
    api_key: Optional[str] = None

class ExportWordRequest(BaseModel):
    base_exam: Optional[Dict[str, Any]] = None
    answer_matrix: Optional[Dict[str, Any]] = None
    school_name: Optional[str] = "TRƯỜNG THCS NGUYỄN DU"
    dept_name: Optional[str] = "TỔ KHOA HỌC TỰ NHIÊN"
    school_year: Optional[str] = "2025 - 2026"
    exam_code: Optional[str] = "101"
    level: Optional[str] = "thcs"

# =========================================================================
# 2. ĐỒNG BỘ 2 CHIỀU SQLITE DATABASE VỚI GOOGLE DRIVE
# =========================================================================
def get_drive_service():
    scopes = ['https://www.googleapis.com/auth/drive']
    try:
        if SERVICE_ACCOUNT_INFO_JSON and SERVICE_ACCOUNT_INFO_JSON.strip():
            info = json.loads(SERVICE_ACCOUNT_INFO_JSON)
            creds = service_account.Credentials.from_service_account_info(info, scopes=scopes)
        elif os.path.exists(SERVICE_ACCOUNT_FILE):
            creds = service_account.Credentials.from_service_account_file(SERVICE_ACCOUNT_FILE, scopes=scopes)
        else:
            return None
        return build('drive', 'v3', credentials=creds)
    except Exception as e:
        print(f"[Drive Auth Error] {e}")
        return None

def get_drive_file_id_by_name(service, folder_id, file_name):
    try:
        query = f"name = '{file_name}' and trashed = false"
        if folder_id:
            query += f" and '{folder_id}' in parents"
        results = service.files().list(q=query, spaces='drive', fields="files(id, name)").execute()
        files = results.get('files', [])
        if files:
            return files[0]['id']
    except Exception as e:
        print(f"[Drive Search File Error]: {e}")
    return None

def restore_db_from_drive():
    service = get_drive_service()
    if not service:
        return False
    try:
        folder_id = GOOGLE_DRIVE_FOLDER_ID.strip() if GOOGLE_DRIVE_FOLDER_ID else None
        file_id = get_drive_file_id_by_name(service, folder_id, DB_DRIVE_FILE_NAME)
        if file_id:
            request = service.files().get_media(fileId=file_id)
            fh = io.BytesIO()
            downloader = MediaIoBaseDownload(fh, request)
            done = False
            while not done:
                _, done = downloader.next_chunk()
            fh.seek(0)
            with open(DB_FILE, "wb") as f:
                f.write(fh.read())
            print(f"[SQLite Persistence] Đã khôi phục thành công {DB_DRIVE_FILE_NAME} từ Google Drive!")
            return True
    except Exception as e:
        print(f"[SQLite Restore Error]: {e}")
    return False

def backup_db_to_drive():
    service = get_drive_service()
    if not service or not os.path.exists(DB_FILE):
        return False
    try:
        folder_id = GOOGLE_DRIVE_FOLDER_ID.strip() if GOOGLE_DRIVE_FOLDER_ID else None
        file_id = get_drive_file_id_by_name(service, folder_id, DB_DRIVE_FILE_NAME)
        media_upload = MediaFileUpload(DB_FILE, mimetype='application/x-sqlite3', resumable=True)
        if file_id:
            service.files().update(fileId=file_id, media_body=media_upload).execute()
            print(f"[SQLite Persistence] Cập nhật backup {DB_DRIVE_FILE_NAME} lên Google Drive.")
        else:
            file_metadata = {'name': DB_DRIVE_FILE_NAME}
            if folder_id:
                file_metadata['parents'] = [folder_id]
            service.files().create(body=file_metadata, media_body=media_upload, fields='id').execute()
            print(f"[SQLite Persistence] Tạo mới backup {DB_DRIVE_FILE_NAME} trên Google Drive.")
        return True
    except Exception as e:
        print(f"[SQLite Backup Error]: {e}")
    return False

# =========================================================================
# 3. KHỞI TẠO SQLITE DATABASE
# =========================================================================
def init_db():
    if not os.path.exists(DB_FILE) or os.path.getsize(DB_FILE) == 0:
        restore_db_from_drive()

    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("""
        CREATE TABLE IF NOT EXISTS questions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            level_stage TEXT DEFAULT 'thcs',
            grade INTEGER NOT NULL,
            subject TEXT DEFAULT 'Toán học',
            topic TEXT NOT NULL,
            lesson TEXT DEFAULT 'Bài học chung',
            level TEXT NOT NULL,
            cognitive_level TEXT DEFAULT 'M2',
            q_type TEXT NOT NULL,
            content TEXT NOT NULL,
            options_json TEXT,
            correct_answer TEXT,
            explanation TEXT,
            has_image INTEGER DEFAULT 0,
            image_base64 TEXT,
            tikz_code TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    try:
        c.execute("ALTER TABLE questions ADD COLUMN lesson TEXT DEFAULT 'Bài học chung'")
    except Exception:
        pass
    try:
        c.execute("ALTER TABLE questions ADD COLUMN cognitive_level TEXT DEFAULT 'M2'")
    except Exception:
        pass
    conn.commit()
    conn.close()

init_db()

def get_db():
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    return conn

# =========================================================================
# 4. RENDER TIKZ 3D QUA KROKI
# =========================================================================
def render_dynamic_tikz_to_base64(tikz_code: str) -> str:
    if not tikz_code or not str(tikz_code).strip():
        return ""
    clean = str(tikz_code).strip()
    if not clean.startswith(r"\begin{tikzpicture}"):
        clean = f"\\begin{{tikzpicture}}\n{clean}\n\\end{{tikzpicture}}"

    full_latex_doc = f"""\\documentclass[tikz,border=3mm]{{standalone}}
\\usepackage{{amsmath,amssymb,pgfplots}}
\\pgfplotsset{{compat=1.18}}
\\usepackage{{tkz-tab}}
\\usetikzlibrary{{arrows.meta,calc,patterns,positioning,intersections,shapes.geometric,3d}}
\\begin{{document}}
{clean}
\\end{{document}}"""

    try:
        res = requests.post(
            "https://kroki.io/tikz/png",
            data=full_latex_doc.encode('utf-8'),
            headers={'Content-Type': 'text/plain; charset=utf-8'},
            timeout=18
        )
        if res.status_code == 200 and res.content:
            return f"data:image/png;base64,{base64.b64encode(res.content).decode('utf-8')}"
    except Exception as e:
        print(f"[TikZ Render Exception]: {e}")
    return ""

def extract_and_convert_embedded_tikz(text: str):
    if not text:
        return text, ""
    match = re.search(r'(\\begin\{tikzpicture\}[\s\S]*?\\end\{tikzpicture\})', text)
    if match:
        raw_tikz = match.group(1)
        img_b64 = render_dynamic_tikz_to_base64(raw_tikz)
        clean_text = text.replace(raw_tikz, "").replace(r"\newline", "\n").strip()
        return clean_text, img_b64
    return text.replace(r"\newline", "\n"), ""

# =========================================================================
# 5. LÀM SẠCH KÝ TỰ & CHỐNG INVALID ESCAPE
# =========================================================================
def clean_and_parse_json(raw_text: str):
    match = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", raw_text)
    content = match.group(1).strip() if match else raw_text.strip()

    content = re.sub(r'\\text\s*\{([^}]*)\}', r'\1', content)
    content = re.sub(r'\\text\s+', r' ', content)
    content = re.sub(r'\\t(?=[a-zA-ZÀ-ỹ\s])', r' ', content)
    content = re.sub(r'\\times', r'×', content)
    content = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', content)

    try:
        return json.loads(content, strict=False)
    except json.JSONDecodeError:
        pass

    content = re.sub(r'\\([bfrnt])(?=[a-zA-Z])', r'\\\\\1', content)

    def sanitize_inner_string(m):
        val = m.group(0)
        inner = val[1:-1]
        inner = inner.replace('\n', '\\n').replace('\r', '\\r')
        fixed = re.sub(r'\\(?!["\\/bfnrtu])', r'\\\\', inner)
        fixed = re.sub(r'\\u(?![0-9a-fA-F]{4})', r'\\\\u', fixed)
        return f'"{fixed}"'

    sanitized = re.sub(r'"(?:\\.|[^"\\])*"', sanitize_inner_string, content)

    try:
        return json.loads(sanitized, strict=False)
    except json.JSONDecodeError:
        try:
            brute_fixed = re.sub(r'\\(?!")', r'\\\\', content)
            return json.loads(brute_fixed, strict=False)
        except Exception as err:
            raise HTTPException(status_code=500, detail=f"Lỗi cú pháp JSON đề thi từ AI: {str(err)}")

def strip_question_prefix(text: str) -> str:
    return re.sub(r'^(?:\[.*?\]\s*)?(?:Câu|Bài)\s*\d+[\s.:-]*', '', text, flags=re.IGNORECASE).strip()

# =========================================================================
# 6. CHUYỂN ĐỔI LATEX SANG OMML WORD EQUATION
# =========================================================================
MML2OMML_XSL = """<?xml version="1.0" encoding="UTF-8"?>
<xsl:stylesheet version="1.0" xmlns:xsl="http://www.w3.org/1999/XSL/Transform"
    xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
    xmlns:mml="http://www.w3.org/1998/Math/MathML">
  <xsl:template match="mml:math"><m:oMath><xsl:apply-templates/></m:oMath></xsl:template>
  <xsl:template match="mml:mrow"><xsl:apply-templates/></xsl:template>
  <xsl:template match="mml:mi|mml:mn|mml:mo|mml:mtext"><m:r><m:t><xsl:value-of select="."/></m:t></m:r></xsl:template>
  <xsl:template match="mml:mfrac"><m:f><m:num><xsl:apply-templates select="*[1]"/></m:num><m:den><xsl:apply-templates select="*[2]"/></m:den></m:f></xsl:template>
  <xsl:template match="mml:msup"><m:sSup><m:e><xsl:apply-templates select="*[1]"/></m:e><m:sup><xsl:apply-templates select="*[2]"/></m:sup></m:sSup></xsl:template>
  <xsl:template match="mml:msub"><m:sSub><m:e><xsl:apply-templates select="*[1]"/></m:e><m:sub><xsl:apply-templates select="*[2]"/></m:sub></m:sSub></xsl:template>
  <xsl:template match="mml:msqrt"><m:rad><m:radPr><m:degHide m:val="on"/></m:radPr><m:deg/><m:e><xsl:apply-templates/></m:e></m:rad></xsl:template>
</xsl:stylesheet>"""

_xslt_doc = ET.fromstring(MML2OMML_XSL.encode('utf-8')) if ET is not None else None
_transform = ET.XSLT(_xslt_doc) if _xslt_doc is not None else None

def convert_latex_to_omml(latex_str: str) -> Optional[str]:
    if not latex_to_mathml or not _transform:
        return None
    try:
        clean_latex = latex_str.strip().replace("\\triangle", "\\Delta ").replace("\\angle", "\\sphericalangle ").replace("^\\circ", "^{\\circ}")
        clean_latex = re.sub(r'\\text\{([^}]+)\}', r'\1', clean_latex)
        clean_latex = re.sub(r'\\t(?=[a-zA-Z\s])', r' ', clean_latex)
        clean_latex = clean_latex.replace("\t", " ")
        mathml_str = latex_to_mathml(clean_latex)
        mml_root = ET.fromstring(mathml_str.encode('utf-8'))
        omml_elem = _transform(mml_root)
        return ET.tostring(omml_elem, encoding='unicode')
    except Exception:
        return None

def add_math_paragraph(doc, text: str, prefix: str = "", font_size: float = 11.0, is_bold: bool = False, left_indent: float = 0, badge_text: str = "", badge_color: str = "blue"):
    p = doc.add_paragraph()
    if left_indent > 0:
        p.paragraph_format.left_indent = Inches(left_indent)
    p.paragraph_format.space_after = Pt(3.5)

    if prefix:
        r_pre = p.add_run(prefix)
        r_pre.bold = is_bold
        r_pre.font.name = "Times New Roman"
        r_pre.font.size = Pt(font_size)

    if badge_text:
        r_badge = p.add_run(f"[{badge_text}] ")
        r_badge.bold = True
        r_badge.font.name = "Times New Roman"
        r_badge.font.size = Pt(font_size - 1.0)
        if badge_color == "blue":
            r_badge.font.color.rgb = RGBColor(37, 99, 235)
        elif badge_color == "green":
            r_badge.font.color.rgb = RGBColor(16, 149, 193)
        elif badge_color == "amber":
            r_badge.font.color.rgb = RGBColor(217, 119, 6)
        elif badge_color == "red":
            r_badge.font.color.rgb = RGBColor(225, 29, 72)

    clean_text = (text or "").replace("\t", " ").replace("\\text", "").replace("\\times", "×")
    clean_text = clean_text.replace("ext{", "").replace("ext", " ")

    tokens = re.split(r'(\$[^$]+\$)', clean_text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('$') and tok.endswith('$') and len(tok) > 2:
            latex_expr = tok[1:-1].strip()
            omml_xml = convert_latex_to_omml(latex_expr)
            if omml_xml and ("<m:oMath" in omml_xml or "<oMath" in omml_xml):
                try:
                    p._p.append(parse_xml(omml_xml.strip()))
                    continue
                except Exception:
                    pass
            sub_text = (
                latex_expr.replace('\\circ', '°')
                .replace('\\cdot', '·')
                .replace('\\pm', '±')
                .replace('\\in', '∈')
                .replace('\\notin', '∉')
                .replace('\\subset', '⊂')
                .replace('\\sqrt', '√')
            )
            sub_text = re.sub(r'\^(\d+)', r'^\1', sub_text)
            r = p.add_run(f" {sub_text} ")
            r.italic = True
            r.font.name = "Cambria Math"
            r.font.size = Pt(font_size)
        else:
            r = p.add_run(tok)
            r.bold = is_bold
            r.font.name = "Times New Roman"
            r.font.size = Pt(font_size)
    return p

# =========================================================================
# 7. ĐỌC ĐỀ MẪU TỪ TỆP MOCK_EXAM_GDPT2018.TEX
# =========================================================================
def load_mock_exam_from_tex(req: GenerateExamRequest, tex_path: str = MOCK_TEX_FILE) -> Dict[str, Any]:
    if not os.path.exists(tex_path):
        raise HTTPException(status_code=404, detail=f"Không tìm thấy tệp đề mẫu {tex_path}")

    with open(tex_path, "r", encoding="utf-8") as f:
        full_content = f.read()

    pattern = rf"\\begin\{{exam\}}\{{([^}}]*)}}\{{{req.grade}\}}\{{{req.level.lower()}\}}\{{(\d+)\}}\{{([^}}]*)\}}([\s\S]*?)\\end\{{exam\}}"
    match = re.search(pattern, full_content)
    content = match.group(6) if match else full_content

    p1_items = []
    p2_items = []
    p3_items = []

    p1_blocks = re.findall(r'\\begin\{question\}\{([M1-4]+)\}([\s\S]*?)\\end\{question\}', content)
    for cog, body in p1_blocks:
        ans_match = re.search(r'\\answer\{([A-D])\}', body)
        answer = ans_match.group(1) if ans_match else "A"
        clean_b = re.sub(r'\\choice\{[A-D]\}\{.*?\}', '', body)
        clean_b = re.sub(r'\\answer\{[A-D]\}', '', clean_b)
        q_text, img_b64 = extract_and_convert_embedded_tikz(clean_b.strip())
        
        choices = re.findall(r'\\choice\{([A-D])\}\{(.*?)\}', body)
        options = [{"id": cid, "text": ctxt} for cid, ctxt in choices]
        p1_items.append({
            "content": q_text,
            "cognitive_level": cog,
            "options": options,
            "answer": answer,
            "image_base64": img_b64,
            "has_image": 1 if img_b64 else 0
        })

    p2_blocks = re.findall(r'\\begin\{questionTF\}\{([M1-4]+)\}([\s\S]*?)\\end\{questionTF\}', content)
    for cog, body in p2_blocks:
        prompt_match = re.search(r'^([\s\S]*?)(?=\\itemTF)', body.strip())
        prompt = prompt_match.group(1).strip() if prompt_match else ""
        expl_match = re.search(r'\\explanation\{(.*?)\}', body)
        explanation = expl_match.group(1) if expl_match else ""
        
        items = re.findall(r'\\itemTF\{([a-d])\}\{(.*?)\}\{(True|False)\}', body)
        statements = {k: v for k, v, _ in items}
        answers = {k: (ans == "True") for k, _, ans in items}
        p2_items.append({
            "prompt": prompt,
            "cognitive_level": cog,
            "statements": statements,
            "answers": answers,
            "explanation": explanation
        })

    p3_blocks = re.findall(r'\\begin\{questionShort\}\{([M1-4]+)\}\{([\d.]+)\}([\s\S]*?)\\end\{questionShort\}', content)
    for cog, pts, body in p3_blocks:
        sol_match = re.search(r'\\solution\{(.*?)\}', body)
        ans_match = re.search(r'\\answer\{(.*?)\}', body)
        solution = sol_match.group(1) if sol_match else ""
        answer = ans_match.group(1) if ans_match else ""
        clean_b = re.sub(r'\\solution\{.*?\}', '', body)
        clean_b = re.sub(r'\\answer\{.*?\}', '', clean_b).strip()
        p3_items.append({
            "content": clean_b,
            "cognitive_level": cog,
            "points": float(pts),
            "solution": solution,
            "answer": answer
        })

    p3_key = "part_3_short_ans" if req.level.lower() == "thpt" else "part_3_essay"
    return {
        "status": "success",
        "source": "tex_template",
        "base_exam": {
            "subject": req.subject,
            "grade": req.grade,
            "level": req.level,
            "part_1_mcq": p1_items,
            "part_2_tf": p2_items,
            p3_key: p3_items
        },
        "answer_matrix": {
            "101": { f"Câu {i+1}": "A" for i in range(len(p1_items)) },
            "102": { f"Câu {i+1}": "B" for i in range(len(p1_items)) },
            "103": { f"Câu {i+1}": "C" for i in range(len(p1_items)) },
            "104": { f"Câu {i+1}": "D" for i in range(len(p1_items)) }
        }
    }

def extract_docx_text_and_images(docx_path: str):
    doc = docx.Document(docx_path)
    text_chunks = []
    images_bytes = []
    for p in doc.paragraphs:
        if p.text.strip():
            text_chunks.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_chunks.append(" | ".join(row_text))
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                images_bytes.append(rel.target_part.blob)
    except Exception as img_err:
        print(f"[Docx Images Warning]: {img_err}")
    return "\n".join(text_chunks), images_bytes

def pdf_to_page_images(pdf_path: str) -> list[bytes]:
    if fitz is None:
        raise HTTPException(status_code=500, detail="PyMuPDF chưa được cài đặt.")
    doc = fitz.open(pdf_path)
    images_bytes = []
    for page in doc:
        pix = page.get_pixmap(dpi=150)
        images_bytes.append(pix.tobytes("jpeg"))
    doc.close()
    return images_bytes

# =========================================================================
# 8. API SINH ĐỀ AI
# =========================================================================
TIKZ_STYLE_GUIDE = """
HƯỚNG DẪN DỰNG HÌNH TIKZ CHUẨN MỰC BẮT BUỘC:
1. Đối với Hình học không gian: BẮT BUỘC dùng nét đứt [dashed] cho các cạnh khuất ở phía sau. Dùng nét liền đậm [thick] cho các cạnh nhìn thấy.
2. Đối với Bảng biến thiên (THPT): Sử dụng môi trường \\begin{tikzpicture} vẽ bảng phân ô, mũi tên biến thiên màu xanh [->, thick, blue].
"""

@app.post("/api/exams/generate-full")
def generate_full_exam(req: GenerateExamRequest):
    if req.selected_question_ids and len(req.selected_question_ids) > 0:
        conn = get_db()
        c = conn.cursor()
        placeholders = ','.join(['?'] * len(req.selected_question_ids))
        c.execute(f"SELECT * FROM questions WHERE id IN ({placeholders})", req.selected_question_ids)
        rows = c.fetchall()
        conn.close()

        p1, p2, p3 = [], [], []
        for r in rows:
            item = dict(r)
            item["options"] = json.loads(item["options_json"]) if item.get("options_json") else []
            if item["q_type"] == "single_choice":
                p1.append(item)
            elif item["q_type"] == "true_false":
                p2.append(item)
            else:
                p3.append(item)

        p3_key = "part_3_short_ans" if req.level.lower() == "thpt" else "part_3_essay"
        return {
            "status": "success",
            "source": "manual_selected_bank",
            "base_exam": {
                "subject": req.subject,
                "grade": req.grade,
                "level": req.level,
                "part_1_mcq": p1,
                "part_2_tf": p2,
                p3_key: p3
            },
            "answer_matrix": {
                "101": { f"Câu {i+1}": "A" for i in range(len(p1)) },
                "102": { f"Câu {i+1}": "B" for i in range(len(p1)) },
                "103": { f"Câu {i+1}": "C" for i in range(len(p1)) },
                "104": { f"Câu {i+1}": "D" for i in range(len(p1)) }
            }
        }

    key_to_use = req.api_key.strip() if req.api_key and req.api_key.strip() else DEFAULT_ENV_KEY
    if not key_to_use:
        return load_mock_exam_from_tex(req)

    caller = genai.Client(api_key=key_to_use)
    level_norm = req.level.lower()
    is_90m = req.duration_minutes >= 90

    if level_norm == "thpt":
        p1_count = 12
        p2_count = 4 if is_90m else 2
        p3_count = 6 if is_90m else 3
        p3_key = "part_3_short_ans"
        p3_desc = f"Phần III: Chính xác {p3_count} câu trắc nghiệm trả lời ngắn (key '{p3_key}')"
    else:
        p1_count = 12 if is_90m else 6
        p2_count = 4 if is_90m else 1
        p3_count = 4 if is_90m else 2
        p3_key = "part_3_essay"
        p3_desc = f"Phần III: Chính xác {p3_count} bài tự luận (key '{p3_key}')"

    total_expected_q = p1_count + p2_count + p3_count
    random_seed = random.randint(100000, 999999)

    system_prompt = f"""Bạn là Chuyên gia Khảo thí GDPT 2018 môn {req.subject} Cấp {req.level.upper()} Lớp {req.grade}.
Nhiệm vụ: Sáng tác mới hoàn toàn đề kiểm tra cho Bài học "{req.lesson}" thuộc Chủ đề "{req.topic}", mức độ: {req.cognitive_level}.
Cấu trúc: Phần I ({p1_count} câu trắc nghiệm), Phần II ({p2_count} câu Đúng/Sai), {p3_desc}.
Mỗi câu có trường "cognitive_level" ("M1", "M2", "M3", "M4").
Mọi dấu '\\' trong LaTeX và TikZ PHẢI ESCAPE THÀNH '\\\\'.
Chỉ xuất JSON thuần túy.
"""

    user_content = f"Lớp {req.grade}, Bài: {req.lesson}, Chủ đề: {req.topic}, Mức độ: {req.cognitive_level}, Thời gian: {req.duration_minutes}p, Mã: {random_seed}."

    try:
        model_name = req.model_name if req.model_name in ["gemini-2.5-flash", "gemini-2.5-pro"] else "gemini-2.5-flash"
        response = caller.models.generate_content(
            model=model_name,
            contents=[system_prompt, user_content],
            config=types.GenerateContentConfig(response_mime_type="application/json", temperature=0.75, max_output_tokens=8192)
        )
        data = clean_and_parse_json(response.text)
        base_ex = data.get("base_exam", data)

        for q in base_ex.get("part_1_mcq", []):
            t_code = q.get("tikz_code")
            if t_code and str(t_code).strip() not in ["null", "None", ""]:
                q["image_base64"] = render_dynamic_tikz_to_base64(t_code)
                q["has_image"] = 1
        for q in base_ex.get("part_2_tf", []):
            t_code = q.get("tikz_code")
            if t_code and str(t_code).strip() not in ["null", "None", ""]:
                q["image_base64"] = render_dynamic_tikz_to_base64(t_code)
                q["has_image"] = 1

        return {
            "status": "success",
            "base_exam": base_ex,
            "answer_matrix": data.get("answer_matrix", {})
        }
    except Exception as e:
        print(f"[Gemini AI Error]: {e}")
        raise HTTPException(status_code=500, detail=f"Lỗi gọi Gemini AI: {str(e)}")

# =========================================================================
# 9. API KHO CÂU HỎI & QUẢN TRỊ
# =========================================================================
@app.get("/api/questions/filter")
def filter_questions(
    level_stage: Optional[str] = None,
    grade: Optional[int] = None,
    q_type: Optional[str] = None,
    lesson: Optional[str] = None,
    cognitive_level: Optional[str] = None,
    keyword: Optional[str] = None
):
    conn = get_db()
    c = conn.cursor()
    query = "SELECT * FROM questions WHERE 1=1"
    params = []
    if level_stage:
        query += " AND level_stage = ?"; params.append(level_stage)
    if grade:
        query += " AND grade = ?"; params.append(grade)
    if q_type:
        query += " AND q_type = ?"; params.append(q_type)
    if lesson and lesson != "Tất cả bài học":
        query += " AND lesson LIKE ?"; params.append(f"%{lesson}%")
    if cognitive_level and cognitive_level != "Tất cả mức độ":
        query += " AND cognitive_level = ?"; params.append(cognitive_level)
    if keyword:
        query += " AND content LIKE ?"; params.append(f"%{keyword}%")

    query += " ORDER BY id DESC LIMIT 150"
    c.execute(query, params)
    rows = c.fetchall()
    conn.close()

    results = []
    for r in rows:
        item = dict(r)
        item["options"] = json.loads(item["options_json"]) if item.get("options_json") else []
        results.append(item)
    return {"total": len(results), "questions": results}

@app.post("/api/questions/save")
def save_question(q: QuestionSave):
    conn = get_db()
    c = conn.cursor()
    c.execute("""
        INSERT INTO questions (
            level_stage, grade, subject, topic, lesson, level, cognitive_level, q_type,
            content, options_json, correct_answer, explanation,
            has_image, image_base64, tikz_code
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        q.level_stage, q.grade, q.subject, q.topic, q.lesson or "Bài học chung",
        q.level, q.cognitive_level or "M2", q.q_type,
        strip_question_prefix(q.content),
        json.dumps(q.options, ensure_ascii=False) if q.options else None,
        q.correct_answer, q.explanation,
        q.has_image, q.image_base64, q.tikz_code
    ))
    new_id = c.lastrowid
    conn.commit()
    conn.close()
    backup_db_to_drive()
    return {"status": "success", "id": new_id}

@app.post("/api/drive/upload-db")
def api_upload_db_to_drive():
    if backup_db_to_drive():
        return {"status": "success", "message": "Đã đồng bộ Kho câu hỏi lên Google Drive thành công!"}
    raise HTTPException(status_code=500, detail="Không thể kết nối Google Drive.")

@app.post("/api/drive/restore-db")
def api_restore_db_from_drive():
    if restore_db_from_drive():
        conn = get_db(); c = conn.cursor()
        c.execute("SELECT COUNT(*) FROM questions"); count = c.fetchone()[0]
        conn.close()
        return {"status": "success", "message": f"Đã khôi phục thành công! Tổng cộng: {count} câu hỏi.", "count": count}
    raise HTTPException(status_code=404, detail="Không tìm thấy exam_bank.db trên Google Drive!")

@app.post("/api/upload-universal")
async def upload_universal(
    file: UploadFile = File(...),
    level_stage: str = Form("thcs"),
    grade: int = Form(7),
    subject: str = Form("Toán học"),
    topic: str = Form("Tổng hợp"),
    api_key: Optional[str] = Form(None)
):
    key_to_use = api_key.strip() if api_key and api_key.strip() else DEFAULT_ENV_KEY
    if not key_to_use:
        raise HTTPException(status_code=400, detail="Chưa có Gemini API Key.")

    caller = genai.Client(api_key=key_to_use)
    filename = file.filename.lower()
    temp_dir = os.path.join(DATA_DIR, "temp_uploads")
    os.makedirs(temp_dir, exist_ok=True)
    temp_path = os.path.join(temp_dir, file.filename)

    with open(temp_path, "wb") as f:
        f.write(await file.read())

    image_parts = []
    text_content = ""
    try:
        if filename.endswith((".docx", ".doc")):
            text_content, embedded_imgs = extract_docx_text_and_images(temp_path)
            for img_b in embedded_imgs[:8]:
                image_parts.append(types.Part.from_bytes(data=img_b, mime_type="image/jpeg"))
        elif filename.endswith(".pdf"):
            for img_b in pdf_to_page_images(temp_path)[:10]:
                image_parts.append(types.Part.from_bytes(data=img_b, mime_type="image/jpeg"))
        else:
            with open(temp_path, "rb") as img_file:
                image_parts.append(types.Part.from_bytes(data=img_file.read(), mime_type=file.content_type or "image/jpeg"))
    finally:
        if os.path.exists(temp_path): os.remove(temp_path)

    prompt = f"Bóc tách từng câu hỏi đề thi môn {subject} cấp {level_stage.upper()} lớp {grade}, chủ đề '{topic}'. Xuất mảng JSON thuần túy [{{...}}]."
    contents = [prompt]
    if text_content: contents.append(f"NỘI DUNG:\n{text_content}")
    contents.extend(image_parts)

    res = caller.models.generate_content(model="gemini-2.5-flash", contents=contents)
    questions = clean_and_parse_json(res.text)
    for q in questions:
        q["level_stage"] = level_stage; q["grade"] = grade; q["subject"] = subject; q["topic"] = topic; q["lesson"] = topic
        q["content"] = strip_question_prefix(q.get("content", ""))
        if q.get("has_image") == 1 and q.get("tikz_code"):
            q["image_base64"] = render_dynamic_tikz_to_base64(q["tikz_code"])
        else:
            q["image_base64"] = ""

    return {"status": "success", "questions": questions}

# =========================================================================
# 10. XUẤT FILE WORD NATIVE (.DOCX)
# =========================================================================
@app.post("/api/exams/export-word")
async def export_word_docx(request: Request):
    try:
        data = await request.json()
    except Exception:
        data = {}

    try:
        base = data.get("base_exam") or {}
        school_name = data.get("school_name") or "SỞ GIÁO DỤC VÀ ĐÀO TẠO"
        dept_name = data.get("dept_name") or "TRƯỜNG HỌC PHỔ THÔNG"
        school_year = data.get("school_year") or "2025 - 2026"
        exam_code = data.get("exam_code") or "101"
        level = (data.get("level") or base.get("level") or "thcs").lower()

        doc = docx.Document()
        for sec in doc.sections:
            sec.top_margin = Inches(0.75); sec.bottom_margin = Inches(0.75)
            sec.left_margin = Inches(0.85); sec.right_margin = Inches(0.75)

        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(3.2)
        table.columns[1].width = Inches(3.8)

        c0 = table.cell(0, 0)
        p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0_1 = p0.add_run(f"{school_name.upper()}\n{dept_name.upper()}\n")
        r0_1.bold = True; r0_1.font.name = "Times New Roman"; r0_1.font.size = Pt(10)
        r0_2 = p0.add_run("\nHọ và tên: .......................................\nSố báo danh: ........... Lớp: ..........")
        r0_2.font.name = "Times New Roman"; r0_2.font.size = Pt(9.5)

        c1 = table.cell(0, 1)
        p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1_1 = p1.add_run(f"{base.get('exam_title', 'ĐỀ KIỂM TRA ĐỊNH KỲ')}\n")
        r1_1.bold = True; r1_1.font.name = "Times New Roman"; r1_1.font.size = Pt(11)
        r1_2 = p1.add_run(f"NĂM HỌC {school_year}\n"); r1_2.font.name = "Times New Roman"; r1_2.font.size = Pt(10)
        r1_3 = p1.add_run(f"Môn: {base.get('subject', 'TOÁN')} {base.get('grade', 12)} - Thời gian: 90 phút\n")
        r1_3.bold = True; r1_3.font.name = "Times New Roman"; r1_3.font.size = Pt(10)
        r1_4 = p1.add_run(f"[ MÃ ĐỀ: {exam_code} ]"); r1_4.bold = True; r1_4.font.name = "Times New Roman"; r1_4.font.size = Pt(10.5)

        p_div = doc.add_paragraph("―" * 48); p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Phần I
        p_p1 = doc.add_paragraph()
        r = p_p1.add_run("PHẦN I. CÂU TRẮC NGHIỆM NHIỀU PHƯƠNG ÁN LỰA CHỌN" + (" (3,0 điểm)" if level == "thpt" else " (4,0 điểm)"))
        r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(11)

        for idx, q in enumerate(base.get('part_1_mcq', [])):
            cog = q.get('cognitive_level', 'M1')
            cog_map = {"M1": ("Nhận biết", "blue"), "M2": ("Thông hiểu", "green"), "M3": ("Vận dụng", "amber"), "M4": ("Vận dụng cao", "red")}
            b_text, b_color = cog_map.get(cog, ("Nhận biết", "blue"))
            add_math_paragraph(doc, q.get('content', ''), prefix=f"Câu {idx+1}: ", font_size=10.5, badge_text=b_text, badge_color=b_color)

            img_b64 = q.get('image_base64')
            if img_b64 and "base64," in str(img_b64):
                try:
                    doc.add_picture(io.BytesIO(base64.b64decode(img_b64.split("base64,")[1])), width=Inches(2.5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass

            raw_opts = q.get('options', [])
            opts_dict = {'A': '', 'B': '', 'C': '', 'D': ''}
            if isinstance(raw_opts, list):
                for i, o in enumerate(raw_opts):
                    lbl = chr(65 + i) if i < 4 else ''
                    opts_dict[lbl] = o.get('text', '') if isinstance(o, dict) else str(o)

            tbl = doc.add_table(rows=1, cols=4); tbl.autofit = False
            for col_idx, key in enumerate(['A', 'B', 'C', 'D']):
                cell = tbl.cell(0, col_idx); cell.width = Inches(1.65)
                p_c = cell.paragraphs[0]; p_c.paragraph_format.space_after = Pt(2)
                r_k = p_c.add_run(f"{key}. "); r_k.bold = True; r_k.font.size = Pt(10.5)
                p_c.add_run(opts_dict.get(key, ''))

        file_stream = io.BytesIO()
        doc.save(file_stream); file_stream.seek(0)
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="De_Thi_{level.upper()}.docx"'}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Lỗi xuất docx: {str(e)}")

@app.get("/")
def read_root():
    return {"status": "online", "service": "EduExam Backend Engine (GDPT 2018 + TeX Template)"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=int(os.environ.get("PORT", 10000)), reload=False)