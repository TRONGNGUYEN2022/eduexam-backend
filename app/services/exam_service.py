import os
import json
import re
import io
from typing import Optional, Dict, Any
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Khởi tạo Gemini Client
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")
client = genai.Client(api_key=GEMINI_API_KEY) if GEMINI_API_KEY else None


def get_system_instruction() -> str:
    return """Bạn là chuyên gia khảo thí và biên soạn đề kiểm tra chuẩn chương trình GDPT 2018 THCS (TP.HCM).

QUY TẮC BẮT BUỘC VỀ ĐỊNH DẠNG CÔNG THỨC TOÁN (LATEX EQUATION):
1. Toàn bộ biểu thức toán học, phép tính, biến số hoặc một dòng biến đổi PHẢI ĐƯỢC BỌC TRONG DUY NHẤT MỘT CẶP DẤU $ (inline) hoặc $$ (display).
   - ĐÚNG: "$A = (x + 2y)(x - 2y) = x^2 - 4y^2 = -2xy - 3y^2$"
   - ĐÚNG: "$\sqrt{4(x-5)} + \frac{1}{3}\sqrt{9(x-5)} = 4$"
   - SAI (NGHIÊM CẤM): "$A = (x+2y)$ $= (x^2)$" hoặc "$A = $ $(x+2y)$" hoặc "$ ... $ $= $".
2. Tuyệt đối KHÔNG để các phép tính (=, +, -, \cdot, \times, \neq, \ge, \le, \in, \Leftrightarrow, \implies) đứng trần ngoài dấu $.
3. Phân tách rõ ràng các bước giải bằng xuống dòng, không dính liền khối.

CẤU TRÚC JSON BẮT BUỘC TRẢ VỀ:
{
  "base_exam": {
    "subject": "Toán học",
    "grade": 6,
    "duration_minutes": 45,
    "exam_title": "ĐỀ KIỂM TRA ĐỊNH KỲ",
    "part_1_mcq": [
      {
        "content": "Nội dung câu hỏi trắc nghiệm kèm công thức $...$",
        "options": [
          {"id": "A", "text": "Phương án A"},
          {"id": "B", "text": "Phương án B"},
          {"id": "C", "text": "Phương án C"},
          {"id": "D", "text": "Phương án D"}
        ]
      }
    ],
    "part_2_tf": [
      {
        "prompt": "Nội dung dẫn dắt câu đúng sai...",
        "statements": {
          "a": "Khẳng định a...",
          "b": "Khẳng định b...",
          "c": "Khẳng định c...",
          "d": "Khẳng định d..."
        },
        "answers": { "a": "Đ", "b": "S", "c": "Đ", "d": "S" },
        "explanation": "Giải thích chi tiết các ý có công thức bọc kín trong $...$"
      }
    ],
    "part_3_essay": [
      {
        "id": 1,
        "points": 1.5,
        "content": "Đề bài tự luận...",
        "solution": "Lời giải chi tiết từng bước có công thức bọc trong $...$",
        "scoring_guide": [
          { "step": "Bước 1: ...", "points": "0.5 đ" },
          { "step": "Bước 2: ...", "points": "0.5 đ" },
          { "step": "Bước 3: ...", "points": "0.5 đ" }
        ]
      }
    ]
  },
  "answer_matrix": {
    "101": { "Câu 1": "A", "Câu 2": "B", "Câu 3": "C", "Câu 4": "D" },
    "102": { "Câu 1": "B", "Câu 2": "C", "Câu 3": "D", "Câu 4": "A" },
    "103": { "Câu 1": "C", "Câu 2": "D", "Câu 3": "A", "Câu 4": "B" },
    "104": { "Câu 1": "D", "Câu 2": "A", "Câu 3": "B", "Câu 4": "C" }
  }
}
"""


def create_full_exam_package(
    subject: str,
    grade: int,
    topic: str,
    duration_minutes: int,
    exam_title: str,
    matrix_spec: Optional[str] = None,
    source_notes: Optional[str] = None,
    essay_mode: Optional[str] = "ai_auto",
    essay_teacher_answer: Optional[str] = None,
    model_name: Optional[str] = "gemini-2.5-flash"
) -> str:
    """Gọi Gemini và trả về NGUYÊN BẢN chuỗi raw JSON"""
    if not client:
        raise ValueError("GEMINI_API_KEY chưa được cấu hình trên server.")

    user_prompt = f"""Hãy biên soạn một đề kiểm tra hoàn chỉnh theo thông số sau:
- Môn học: {subject}
- Khối lớp: Lớp {grade}
- Thời gian làm bài: {duration_minutes} phút
- Kỳ kiểm tra: {exam_title}
- Chủ đề kiến thức: {topic}
"""

    if matrix_spec:
        user_prompt += f"\n- KHUNG MA TRẬN ĐẶC TẢ YÊU CẦU: {matrix_spec}"
    if source_notes:
        user_prompt += f"\n- TÀI LIỆU NGUỒN / CÂU HỎI THAM KHẢO: {source_notes}"
    if essay_mode in ["teacher_only", "hybrid"] and essay_teacher_answer:
        user_prompt += f"\n- ĐÁP ÁN NẠP VÀO TỪ GIÁO VIÊN: {essay_teacher_answer}"

    user_prompt += "\nTrả về DUY NHẤT 1 đối tượng JSON nguyên vẹn theo đúng cấu trúc schema đã quy định."

    target_model = model_name or "gemini-2.5-flash"

    response = client.models.generate_content(
        model=target_model,
        contents=user_prompt,
        config=types.GenerateContentConfig(
            system_instruction=get_system_instruction(),
            response_mime_type="application/json"
        )
    )

    raw_text = response.text or "{}"
    return raw_text


def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    """Thiết lập khoảng đệm lề cho cell trong bảng word"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def generate_docx_file(data: dict) -> io.BytesIO:
    """Hàm tạo file Word .docx chuyên nghiệp, căn chỉnh đều đặn bằng bảng ẩn cho trắc nghiệm"""
    doc = Document()
    
    # Thiết lập lề trang giấy A4
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(0.8)

    # Thiết lập font chữ mặc định Times New Roman
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    font.color.rgb = RGBColor(0, 0, 0)

    base_exam = data.get("base_exam", {})
    matrix = data.get("answer_matrix", {})
    school_name = data.get("school_name", "TRƯỜNG THCS LÊ QUÝ ĐÔN (QUẬN 3)")
    dept_name = data.get("dept_name", "TỔ KHOA HỌC TỰ NHIÊN")
    school_year = data.get("school_year", "2026 - 2027")
    exam_code = data.get("exam_code", "101")

    # --- HEADER ĐỀ THI 2 CỘT ---
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False

    col_widths = [Inches(3.2), Inches(3.3)]
    for row in header_table.rows:
        for i, width in enumerate(col_widths):
            row.cells[i].width = width

    left_cell = header_table.cell(0, 0)
    p_left_1 = left_cell.paragraphs[0]
    p_left_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p_left_1.add_run(school_name)
    run1.bold = True
    run1.font.size = Pt(11)

    p_left_2 = left_cell.add_paragraph()
    p_left_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p_left_2.add_run(dept_name)
    run2.bold = True
    run2.font.size = Pt(11)

    p_left_3 = left_cell.add_paragraph()
    p_left_3.paragraph_format.space_before = Pt(4)
    p_left_3.add_run("Họ và tên: ....................................................\nSố báo danh: ................ Lớp: ................")

    right_cell = header_table.cell(0, 1)
    p_right_1 = right_cell.paragraphs[0]
    p_right_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_right_1.add_run(base_exam.get("exam_title", "ĐỀ KIỂM TRA ĐỊNH KỲ"))
    r_title.bold = True
    r_title.font.size = Pt(12)

    p_right_2 = right_cell.add_paragraph()
    p_right_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_right_2.add_run(f"NĂM HỌC {school_year}\nMôn: {base_exam.get('subject', 'Toán')} {base_exam.get('grade', 6)} - Thời gian: {base_exam.get('duration_minutes', 45)} phút")

    p_right_3 = right_cell.add_paragraph()
    p_right_3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_right_3.paragraph_format.space_before = Pt(4)
    run_code = p_right_3.add_run(f"MÃ ĐỀ: {exam_code}")
    run_code.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- PHẦN I: TRẮC NGHIỆM (SỬ DỤNG BẢNG ĐỂ CĂN ĐỀU CÁC CỘT A, B, C, D) ---
    p_p1 = doc.add_paragraph()
    r_p1 = p_p1.add_run("PHẦN I. CÂU TRẮC NGHIỆM NHIỀU PHƯƠNG ÁN LỰA CHỌN (4,0 điểm)")
    r_p1.bold = True

    p_note1 = doc.add_paragraph()
    r_n1 = p_note1.add_run("Thí sinh trả lời từ câu 1 đến hết. Mỗi câu hỏi chỉ chọn một phương án đúng nhất.")
    r_n1.italic = True
    r_n1.font.size = Pt(11)

    p1_list = base_exam.get("part_1_mcq", [])
    for idx, q in enumerate(p1_list):
        doc.add_paragraph(f"Câu {idx + 1}: {q.get('content', '')}")
        
        opts = q.get("options", [])
        opt_table = doc.add_table(rows=2, cols=2)
        opt_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        opt_table.autofit = False
        
        for r_idx in range(2):
            for c_idx in range(2):
                cell = opt_table.cell(r_idx, c_idx)
                cell.width = Inches(3.25)
                set_cell_margins(cell, top=40, bottom=40, left=80, right=80)
                
        opt_pairs = [(0,0,0), (0,1,1), (1,0,2), (1,1,3)]
        for r, c, opt_idx in opt_pairs:
            if opt_idx < len(opts):
                o = opts[opt_idx]
                oid = o.get('id', 'A')
                otext = o.get('text', o.get('content', ''))
                cell = opt_table.cell(r, c)
                cell.text = f"{oid}. {otext}"
        
        doc.add_paragraph().paragraph_format.space_after = Pt(3)

    # --- PHẦN II: ĐÚNG / SAI ---
    p_p2 = doc.add_paragraph()
    p_p2.paragraph_format.space_before = Pt(8)
    r_p2 = p_p2.add_run("PHẦN II. CÂU TRẮC NGHIỆM ĐÚNG / SAI (3,0 điểm)")
    r_p2.bold = True

    p2_list = base_exam.get("part_2_tf", [])
    for idx, q in enumerate(p2_list):
        doc.add_paragraph(f"Câu {idx + 1}: {q.get('prompt', '')}")
        st = q.get("statements", {})
        for k in ['a', 'b', 'c', 'd']:
            if k in st or k.upper() in st:
                val = st.get(k, st.get(k.upper(), ''))
                doc.add_paragraph(f"   {k}) {val}    [Đ]   [S]")

    # --- PHẦN III: TỰ LUẬN ---
    p_p3 = doc.add_paragraph()
    p_p3.paragraph_format.space_before = Pt(8)
    r_p3 = p_p3.add_run("PHẦN III. TỰ LUẬN (3,0 điểm)")
    r_p3.bold = True

    p3_list = base_exam.get("part_3_essay", [])
    for idx, q in enumerate(p3_list):
        pts = q.get("points", 1.5)
        doc.add_paragraph(f"Bài {idx + 1} ({pts} điểm): {q.get('content', '')}")
        for _ in range(3):
            doc.add_paragraph("...................................................................................................................................")

    # --- KẾT THÚC ---
    doc.add_paragraph().paragraph_format.space_before = Pt(10)
    p_end = doc.add_paragraph("--- HẾT ĐỀ THI ---")
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_end.runs[0].bold = True

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream